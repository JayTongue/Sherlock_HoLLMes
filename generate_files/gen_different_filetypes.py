import random
from pathlib import Path
from reportlab.lib.pagesizes import letter
from docx import Document

def write_txt(path: Path, text: str):
    data = text.encode("ascii", errors="ignore")
    if not data:
        data = b"\n"
    with open(path, "wb") as f:
        f.write(data)


def write_rtf(path: Path, text: str):
    safe = (
        text.replace("\\", r"\\")
            .replace("{", r"\{")
            .replace("}", r"\}")
            .replace("\n", r"\par\n")
    )

    header = r"{\rtf1\ansi\deff0" + "\n"
    footer = "}\n"

    with open(path, "wb") as f:
        f.write(header.encode("ascii"))
        f.write(safe.encode("ascii", errors="ignore") or b"\\par\n")
        f.write(footer.encode("ascii"))


def write_pdf(path: Path, text: str):
    from reportlab.lib.units import inch
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    from reportlab.lib.styles import getSampleStyleSheet
    
    doc = SimpleDocTemplate(str(path), pagesize=letter,
                           leftMargin=72, rightMargin=72,
                           topMargin=72, bottomMargin=72)
    
    styles = getSampleStyleSheet()
    story = []
    
    for para in text.split('\n\n'):
        if para.strip():
            safe_para = (
                para.replace('&', '&amp;')
                    .replace('<', '&lt;')
                    .replace('>', '&gt;')
                    .replace('\n', '<br />')
            )
            story.append(Paragraph(safe_para, styles['Normal']))
            story.append(Spacer(1, 12))
    doc.build(story)



def write_docx(path: Path, text: str):
    def sanitize_for_xml(text: str) -> str:
        def is_valid_xml_char(c):
            codepoint = ord(c)
            return (
                codepoint == 0x9 or
                codepoint == 0xA or
                codepoint == 0xD or
                (0x20 <= codepoint <= 0xD7FF) or
                (0xE000 <= codepoint <= 0xFFFD) or
                (0x10000 <= codepoint <= 0x10FFFF)
            )
        return ''.join(c for c in text if is_valid_xml_char(c))
    
    text = sanitize_for_xml(text)
    paragraphs = text.split("\n\n") if text else [""]
    
    doc = Document()
    doc.add_heading("Document", level=1)
    for p in paragraphs:
        if p.strip():
            doc.add_paragraph(p)
    doc.save(path)


def write_htm(path: Path, text: str):
    safe = (
        text.replace("&", "&amp;")
            .replace("<", "&lt;")
            .replace(">", "&gt;")
    )

    html = (
        "<!doctype html>\n"
        "<html><head><meta charset=\"utf-8\"><title>Document</title></head>\n"
        "<body><pre>\n"
        f"{safe}"
        "\n</pre></body></html>\n"
    )

    with open(path, "wb") as f:
        f.write(html.encode("utf-8"))


def write_file(path: Path, ext: str, text: str):
    if ext == "txt":
        write_txt(path, text)
    elif ext == "rtf":
        write_rtf(path, text)
    elif ext == "pdf":
        write_pdf(path, text)
    elif ext == "docx":
        write_docx(path, text)
    elif ext == "htm":
        write_htm(path, text)
    else:
        raise ValueError(f"Unsupported file type: {ext}")
    
    
def gen_different_filetypes_main(path, input_string):
    ft_dict = {
                'pdf': 1, 
                'txt': 1, 
                # 'rtf': 1, 
                'docx': 1, 
                # 'htm': 1
                }
    filetypes = list(ft_dict.keys())
    filetype_weights = list(ft_dict.values())
    # random.seed(34)
    ext = random.choices(filetypes, weights=filetype_weights)[0]
    output_path = Path(f'{str(path).split('.')[0]}.{ext}')
    # print(output_path)
    write_file(output_path, ext, input_string)